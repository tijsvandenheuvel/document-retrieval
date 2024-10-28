import os
from llama_index import GPTSimpleVectorIndex, Document
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

# Step 1: Define helper functions to extract text from files
def extract_text_from_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_word(file_path):
    doc = DocxDocument(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def load_documents_from_directory(directory_path):
    documents = []
    for filename in os.listdir(directory_path):
        file_path = os.path.join(directory_path, filename)
        if filename.endswith(".pdf"):
            text = extract_text_from_pdf(file_path)
        elif filename.endswith(".docx"):
            text = extract_text_from_word(file_path)
        else:
            continue
        documents.append(Document(text, metadata={"filename": filename}))
    return documents

# Set up document directory path
directory_path = "documents"

# Step 2: Check if index exists; if not, create it
index_file = "my_document_index.json"
if os.path.exists(index_file):
    index = GPTSimpleVectorIndex.load_from_disk(index_file)
else:
    documents = load_documents_from_directory(directory_path)
    index = GPTSimpleVectorIndex(documents)
    index.save_to_disk(index_file)

# Step 3: Define a function to search with natural language queries
def search_documents(query):
    response = index.query(query)
    return response

# Step 4: Run a sample query
query = "Give me documents about privacy and security cameras"
results = search_documents(query)

# Display results
print("Search Results:")
for result in results:
    print("Document:", result.metadata["filename"])
    print("Excerpt:", result.text[:200], "...")  # Print the first 200 characters as an excerpt
    print("-" * 50)