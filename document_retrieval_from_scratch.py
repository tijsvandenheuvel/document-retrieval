import os
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

class Document:
    def __init__(self, text, metadata):
        self.text = text
        self.metadata = metadata or {}

def extract_text_from_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_word(file_path):
    doc = DocxDocument(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def load_documents_from_directory(directory_path):
    documents = []
    for filename in os.listdir(directory_path):
        file_path = os.path.join(directory_path, filename)
        if filename.endswith(".pdf"):
            text = extract_text_from_pdf(file_path)
        elif filename.endswith(".docx"):
            text = extract_text_from_word(file_path)
        else:
            continue
        documents.append(Document(text, metadata={"filename": filename}))
    return documents

# Set up document directory path
directory_path = "documents"

corpus_of_documents = load_documents_from_directory(directory_path)

def jaccard_similarity(query, document):
    query = query.lower().split(" ")
    document = document.lower().split(" ")
    intersection = set(query).intersection(set(document))
    union = set(query).union(set(document))
    return len(intersection)/len(union)

def return_best_match(query, corpus):
    similarities = []
    for doc in corpus:
        similarity = jaccard_similarity(query, doc.text)  # Use doc.text for comparison
        similarities.append(similarity)
    best_match_index = similarities.index(max(similarities))
    best_match_document = corpus[best_match_index]
    return best_match_document  # Return the Document object

def return_matching_documents(query, corpus, threshold=0.1):
    matching_documents = []
    for doc in corpus:
        similarity = jaccard_similarity(query, doc.text)
        if similarity >= threshold:
            matching_documents.append((doc, similarity))
    return matching_documents

user_prompt = "What is the document you are looking for about?"

user_input = input(user_prompt)

# get best match

response = return_best_match(user_input, corpus_of_documents)

print(f"The document that best matches the query is: '{response.metadata['filename']}'")
print("Document text excerpt:")
print(response.text[:200]) 

# get other matches

similarity_threshold = 0.1  

matching_documents = return_matching_documents(user_input, corpus_of_documents, similarity_threshold)

if matching_documents:
    print("The following documents also match your query:")
    for doc, similarity in matching_documents:
        print(f"Filename: {doc.metadata['filename']}, Similarity: {similarity:.2f}")
        print(f"Excerpt: {doc.text[:200]}\n")
else:
    print("No documents matched your query.")